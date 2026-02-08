import os
import sqlite3
from datetime import datetime
from io import BytesIO
from typing import List, Dict

from flask import (
    Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
)

from docx import Document

# =====================================================
# CONFIG
# =====================================================
APP_NAME = "EthosJus"
app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-ethosjus-secret-final-v1")

DATA_DIR = os.path.abspath("./data")
os.makedirs(DATA_DIR, exist_ok=True)
DB_PATH = os.path.join(DATA_DIR, "ethosjus.sqlite3")

# =====================================================
# LINKS OFICIAIS (OAB)
# =====================================================
LINKS_OFICIAIS = {
    "codigo_etica_oab_pdf": "https://www.oab.org.br/content/pdf/legislacaoOab/codigodeetica.pdf",
    "provimento_205_2021": "https://www.oab.org.br/leisnormas/legislacao/provimentos/205-2021",
    "estatuto_oab_8906": "https://www.planalto.gov.br/ccivil_03/leis/l8906.htm",
}

# =====================================================
# PERGUNTAS RÁPIDAS (EXPANDIDA)
# =====================================================
QUICK_QUESTIONS: List[Dict[str, str]] = [
    # --- Publicidade / Marketing Digital ---
    {"text": "Posso impulsionar post no Instagram?", "tag": "Publicidade"},
    {"text": "Posso divulgar valores e promoções?", "tag": "Publicidade"},
    {"text": "Posso prometer resultado ou usar 'garantia'?", "tag": "Publicidade"},
    {"text": "Posso postar fotos com clientes ou processos?", "tag": "Publicidade"},
    {"text": "Posso anunciar 'especialista'?", "tag": "Publicidade"},
    {"text": "Posso responder caixinha de perguntas com caso real?", "tag": "Publicidade"},
    {"text": "Posso fazer sorteio de brindes ou serviços?", "tag": "Publicidade"},
    {"text": "Posso usar Google Ads (Links Patrocinados)?", "tag": "Publicidade"},
    {"text": "Posso enviar e-mail marketing ou mala direta?", "tag": "Publicidade"},
    {"text": "Posso usar logotipos de Tribunais no meu cartão?", "tag": "Publicidade"},

    # --- Sigilo / Dados / LGPD ---
    {"text": "Posso falar do caso com familiares do cliente?", "tag": "Sigilo"},
    {"text": "Posso confirmar que a pessoa é minha cliente?", "tag": "Sigilo"},
    {"text": "Como lidar com documentos sensíveis e LGPD?", "tag": "LGPD"},
    {"text": "Posso gravar reunião com cliente?", "tag": "LGPD"},

    # --- Honorários / Financeiro ---
    {"text": "Preciso de contrato de honorários por escrito?", "tag": "Honorários"},
    {"text": "Posso cobrar consulta? Como formalizar?", "tag": "Honorários"},
    {"text": "Como combinar êxito (quota litis) sem abusos?", "tag": "Honorários"},
    {"text": "O que fazer com inadimplência sem expor o cliente?", "tag": "Honorários"},
    {"text": "Posso reter documentos por falta de pagamento?", "tag": "Honorários"},
    {"text": "Posso cobrar abaixo da tabela da OAB?", "tag": "Honorários"},
    {"text": "Posso aceitar bens como pagamento?", "tag": "Honorários"},

    # --- Conflito / Ética / Sociedade ---
    {"text": "Posso atuar contra ex-cliente?", "tag": "Conflito"},
    {"text": "Posso representar duas partes relacionadas?", "tag": "Conflito"},
    {"text": "Quando devo recusar patrocínio?", "tag": "Ética"},
    {"text": "Posso substabelecer sem avisar?", "tag": "Ética"},
    {"text": "Posso ter sociedade com contador ou médico?", "tag": "Sociedade"},
    {"text": "Advogado pode ser preposto do cliente?", "tag": "Ética"},

    # --- Prerrogativas / Atuação ---
    {"text": "Posso falar mal de outro advogado publicamente?", "tag": "Postura"},
    {"text": "Como agir em audiência com urbanidade?", "tag": "Postura"},
    {"text": "Posso atuar sem procuração em urgência?", "tag": "Prerrogativa"},
    {"text": "O que fazer se o cliente 'sumir'?", "tag": "Gestão"},
]

# =====================================================
# HELPERS DE RESPOSTA (HTML)
# =====================================================
def _html_escape(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _make_answer(title: str, bullets: List[str], delicate: bool = True) -> str:
    warn = ""
    if delicate:
        warn = """
        <div class="alert-box warning">
          <strong>Nota ética:</strong> isto é um guia informacional. Em caso concreto, consulte o TED/OAB e a normativa aplicável.
        </div>
        """
    lis = "".join([f"<li>{_html_escape(b)}</li>" for b in bullets if (b or "").strip()])
    return f"""
    <div class="resposta-humanizada">
      <h3>{_html_escape(title)}</h3>
      {warn}
      <ul>{lis}</ul>
    </div>
    """

# =====================================================
# RESPOSTAS PRONTAS (DATABASE)
# =====================================================
RESPOSTAS_DB: Dict[str, str] = {
    # --- Publicidade ---
    "Posso impulsionar post no Instagram?": _make_answer(
        "Pode, com cuidado (Provimento 205/2021).",
        [
            "Em geral, é permitido impulsionar conteúdo informativo, sem oferta direta de serviços.",
            "Evite promessas, comparações, autopromoção agressiva e captação indevida.",
            "Priorize conteúdo educativo (direitos, prazos, orientações gerais) sem chamadas do tipo ‘contrate agora’.",
            "Quando em dúvida, consulte o Provimento 205/2021 e orientações do TED da sua seccional."
        ]
    ),
    "Posso divulgar valores e promoções?": _make_answer(
        "Regra prática: evite apelo comercial.",
        [
            "“Promoção”, ‘desconto’, ‘pacote’ e linguagem mercantil tendem a ser problemáticos.",
            "Se precisar informar valores, prefira informar em contato privado e com sobriedade.",
            "Evite comparações (‘mais barato’, ‘melhor do que’)."
        ]
    ),
    "Posso prometer resultado ou usar 'garantia'?": _make_answer(
        "Não é recomendado — risco ético alto.",
        [
            "Promessa de resultado pode configurar publicidade irregular e ferir deveres de moderação.",
            "Use linguagem de meios, não de fins: explique etapas, riscos e variáveis do caso.",
            "Evite frases absolutas (‘ganho certo’, ‘causa ganha’)."
        ]
    ),
    "Posso postar fotos com clientes ou processos?": _make_answer(
        "Só com extrema cautela — e em muitos casos é melhor evitar.",
        [
            "Pode violar sigilo, privacidade e gerar captação indevida.",
            "Evite prints, nomes, números de processo, documentos, peças e decisões com elementos identificáveis.",
            "Prefira conteúdo genérico: ‘tese X’, ‘tema Y’, sem caso real."
        ]
    ),
    "Posso anunciar 'especialista'?": _make_answer(
        "Use apenas se houver titulação/critério compatível e comunicação sóbria.",
        [
            "Evite títulos chamativos e qualificações vagas (‘o melhor’, ‘o mais renomado’).",
            "Prefira: área de atuação e formação real, sem induzir o público a erro."
        ]
    ),
    "Posso responder caixinha de perguntas com caso real?": _make_answer(
        "Evite. Transforme em exemplo abstrato.",
        [
            "Mesmo sem nome, detalhes podem identificar a pessoa.",
            "Responda em tese: explique regras gerais, limites e caminhos típicos.",
            "Inclua aviso: não é consulta; caso concreto exige análise."
        ]
    ),
    "Posso fazer sorteio de brindes ou serviços?": _make_answer(
        "Não. É vedado expressamente.",
        [
            "A advocacia não pode ser mercantilizada.",
            "Sorteios, brindes e oferta de serviços gratuitos para captar clientela tendem a ser infrações éticas.",
            "O foco deve ser conteúdo informativo e competência técnica."
        ]
    ),
    "Posco usar Google Ads (Links Patrocinados)?": _make_answer(
        "Em regra, é admitido com moderação (Provimento 205/2021).",
        [
            "Mantenha caráter informativo e linguagem sóbria.",
            "Evite ‘consulta grátis’, ‘melhor preço’ e promessas.",
            "Atenção à captação indevida e mercantilização."
        ]
    ),
    "Posso enviar e-mail marketing ou mala direta?": _make_answer(
        "Somente com consentimento e para base própria.",
        [
            "Evite disparos para listas desconhecidas (risco de spam/captação).",
            "Prefira boletins informativos para contatos que autorizaram.",
            "Inclua possibilidade de descadastro."
        ]
    ),
    "Posso usar logotipos de Tribunais no meu cartão?": _make_answer(
        "Não. Evite símbolos oficiais.",
        [
            "Pode induzir a erro sobre vínculo com órgão público.",
            "Use apenas identidade visual própria.",
        ]
    ),

    # --- Sigilo/LGPD ---
    "Posso falar do caso com familiares do cliente?": _make_answer(
        "Não, sem autorização expressa e limites claros.",
        [
            "A regra é confidencialidade.",
            "Se o cliente autorizar, delimite: quem, assunto e finalidade.",
            "Compartilhe o mínimo necessário."
        ]
    ),
    "Posco confirmar que a pessoa é minha cliente?": _make_answer(
        "Evite — o próprio vínculo pode ser sensível.",
        [
            "Resposta padrão segura: ‘Não posso confirmar nem negar informações de atendimento/contratação’.",
            "Exceções devem ser justificadas e, quando possível, autorizadas por escrito."
        ]
    ),
    "Como lidar com documentos sensíveis e LGPD?": _make_answer(
        "Mínimo necessário + controle de acesso.",
        [
            "Guarde só o necessário para o serviço.",
            "Senha forte, 2FA, backup e descarte seguro.",
            "Defina política de acesso e retenção."
        ]
    ),
    "Posco gravar reunião com cliente?": _make_answer(
        "Boa prática: só com consentimento e finalidade definida.",
        [
            "Explique motivo e onde ficará armazenado.",
            "Evite gravação por padrão; prefira ata.",
            "Se o cliente não quiser, não grave."
        ]
    ),

    # --- Honorários ---
    "Preciso de contrato de honorários por escrito?": _make_answer(
        "Altamente recomendado.",
        [
            "Defina escopo, honorários, despesas, pagamentos e rescisão.",
            "Deixe claro o que é extra (recursos, diligências).",
            "Guarde assinado (inclusive eletrônico)."
        ]
    ),
    "Posco cobrar consulta? Como formalizar?": _make_answer(
        "Pode — e registre por escrito.",
        [
            "Informe valor e o que será entregue.",
            "Registre por WhatsApp/e-mail e formalize se virar patrocínio.",
            "Evite promessas de resultado."
        ]
    ),
    "Como combinar êxito (quota litis) sem abusos?": _make_answer(
        "Transparência e moderação.",
        [
            "Explique base de cálculo e quando incide.",
            "Evite percentuais desproporcionais.",
            "Deixe claro custas e sucumbência."
        ]
    ),
    "O que fazer com inadimplência sem expor o cliente?": _make_answer(
        "Negocie e documente; sem exposição.",
        [
            "Tente parcelar e ajustar datas.",
            "Formalize encerramento se necessário.",
            "Proteja prazos e entregue documentos essenciais."
        ]
    ),
    "Posco reter documentos por falta de pagamento?": _make_answer(
        "Evite — risco ético alto.",
        [
            "Cobrança deve ser feita por meios próprios, sem coação.",
            "Em dúvida, consulte o TED/OAB."
        ]
    ),
    "Posco cobrar abaixo da tabela da OAB?": _make_answer(
        "Cuidado com aviltamento.",
        [
            "Valores irrisórios podem caracterizar aviltamento e captação.",
            "Pro bono tem regras e não pode ser usado como publicidade.",
            "Mantenha dignidade e justificativa."
        ]
    ),
    "Posco aceitar bens como pagamento?": _make_answer(
        "Em regra, sim — com cautela.",
        [
            "Registre no contrato e avalie compatibilidade do valor.",
            "Evite vulnerabilidade/lesão do cliente.",
        ]
    ),
}

def generate_answer_for_question(q: str) -> str:
    # Aceita pergunta do front mesmo que venha com espaços extras
    q = (q or "").strip()
    if q in RESPOSTAS_DB:
        return RESPOSTAS_DB[q]
    return _make_answer(
        "Guia ético (resposta geral)",
        [
            "Essa dúvida depende do contexto e da normativa aplicável.",
            "Use a regra do ‘mínimo necessário’, moderação na comunicação e proteção de confidencialidade.",
            "Quando houver risco ético, consulte o TED/OAB e a normativa aplicável."
        ],
        delicate=True
    )

# =====================================================
# DB (SQLITE)
# =====================================================
def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS qa_history (
            id INTEGER PRIMARY KEY,
            question TEXT,
            answer TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()

def save_history(question: str, answer: str):
    conn = db()
    conn.execute(
        "INSERT INTO qa_history (question, answer, created_at) VALUES (?,?,?)",
        (question, answer, datetime.now().strftime("%d/%m %H:%M"))
    )
    conn.commit()
    conn.close()

def get_history(limit: int = 50):
    conn = db()
    rows = conn.execute(
        "SELECT * FROM qa_history ORDER BY id DESC LIMIT ?",
        (limit,)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# =====================================================
# CONTRATO (PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS)
# =====================================================
def gerar_contrato_advocacia(data: dict) -> str:
    contratante = (data.get("contratante") or "CONTRATANTE").strip()
    contratado = (data.get("contratado") or "CONTRATADO(A)").strip()
    oab = (data.get("oab") or "OAB/UF XXXXX").strip()
    foro = (data.get("foro") or "Comarca de __________________/UF").strip()

    objeto = (data.get("objeto") or "Prestação de serviços advocatícios no tema: ____________________________.").strip()
    honorarios = (data.get("honorarios") or "Honorários: R$ ________ (fixo) e/ou ________% (êxito), conforme condições abaixo.").strip()
    despesas = (data.get("despesas") or "Custas, emolumentos e despesas correrão por conta do(a) CONTRATANTE, mediante prestação de contas.").strip()
    comunicacao = (data.get("comunicacao") or "WhatsApp/E-mail").strip()
    rescisao_extra = (data.get("rescisao") or "").strip()

    rescisao_base = """RESCISÃO, RENÚNCIA E ENCERRAMENTO
1. Qualquer das partes poderá rescindir este contrato, mediante comunicação por escrito.
2. Em caso de rescisão pelo(a) CONTRATANTE, serão devidos os honorários proporcionais ao trabalho já realizado, além de despesas comprovadas.
3. Em caso de renúncia pelo(a) CONTRATADO(A), serão adotadas as providências necessárias para evitar prejuízo ao(à) CONTRATANTE, incluindo comunicação formal, entrega de documentos essenciais e orientações de transição, respeitados os prazos e deveres profissionais.
"""
    if rescisao_extra:
        rescisao_base += f"\nCláusula adicional informada pelas partes:\n- {rescisao_extra}\n"

    return f"""CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS (MODELO)

PARTES
CONTRATANTE: {contratante}
CONTRATADO(A): {contratado} — {oab}

1) OBJETO
1.1. O presente contrato tem por objeto {objeto}
1.2. O serviço observa a legislação aplicável, o Código de Ética e Disciplina e a independência técnica do(a) CONTRATADO(A).

2) ESCOPO, LIMITES E ATOS INCLUÍDOS
2.1. Inclui-se, em regra, conforme a natureza do serviço:
(a) reunião/consulta inicial e definição de estratégia;
(b) análise de documentos fornecidos;
(c) elaboração de peças e manifestações necessárias ao objeto;
(d) acompanhamento do andamento e comunicação de eventos relevantes.
2.2. NÃO estão incluídos, salvo ajuste escrito específico:
(a) propositura de novas demandas não previstas no objeto;
(b) recursos em qualquer instância/tribunal;
(c) sustentações orais, memoriais, despachos presenciais;
(d) diligências externas, viagens, audiências extras ou incidentes não previstos;
(e) perícias/assistência técnica especializada fora do escopo.
2.3. Caso surjam medidas não previstas, as partes poderão firmar aditivo com novo escopo e honorários.

3) DEVERES DO(A) CONTRATANTE
3.1. Fornecer informações verdadeiras, completas e documentos necessários, respondendo por omissões que possam comprometer a atuação.
3.2. Manter canais de contato atualizados e atender solicitações em prazo compatível com urgências e prazos.
3.3. Realizar pagamentos pactuados e reembolsar despesas, conforme previsto.

4) HONORÁRIOS
4.1. As partes ajustam: {honorarios}
4.2. Honorários podem ser fixos, por fase/etapa ou por êxito (quando aplicável).
4.3. Honorários de êxito (se aplicáveis):
(a) incidem sobre o benefício econômico efetivamente obtido pelo(a) CONTRATANTE;
(b) são devidos em acordo, sentença, recebimento administrativo, compensação ou forma equivalente;
(c) se houver acordo sem participação do(a) CONTRATADO(A), poderão ser devidos conforme atuação já realizada, conforme pactuação.
4.4. Honorários de sucumbência:
(a) quando fixados em favor do(a) advogado(a), pertencem ao(à) CONTRATADO(A), sem prejuízo dos honorários contratuais, salvo ajuste expresso em contrário.
4.5. Atraso e inadimplência:
(a) a falta de pagamento autoriza suspensão de atos não urgentes, com comunicação ao(à) CONTRATANTE;
(b) cobranças devem observar urbanidade e discrição, sem exposição.

5) DESPESAS, CUSTAS E REEMBOLSO
5.1. {despesas}
5.2. Despesas incluem: custas, emolumentos, diligências, cópias, autenticações, deslocamentos, correspondentes e taxas.
5.3. Sempre que possível, o(a) CONTRATADO(A) informará previamente despesas relevantes. Em urgência, poderão ser realizadas para evitar prejuízo, com posterior prestação de contas.

6) COMUNICAÇÃO E ATUALIZAÇÕES
6.1. Canal preferencial: {comunicacao}
6.2. O(a) CONTRATADO(A) comunicará eventos relevantes e necessidades de documentos.
6.3. Mensagens são voltadas à logística e atualizações; análises extensas serão priorizadas em reunião/consulta.

7) CONFIDENCIALIDADE E PROTEÇÃO DE DADOS
7.1. As partes se comprometem a manter sigilo sobre informações e documentos relacionados ao caso.
7.2. Dados e documentos serão tratados estritamente para execução do contrato e cumprimento de deveres profissionais/legais.
7.3. Adotam-se medidas razoáveis de segurança (controle de acesso, armazenamento adequado e descarte seguro).

8) INDEPENDÊNCIA TÉCNICA
8.1. O(a) CONTRATADO(A) atuará com independência técnica, não se comprometendo com resultado específico.

9) PRAZO
9.1. Vigência a partir da assinatura até a conclusão do objeto, rescisão ou encerramento.

10) {rescisao_base}

11) FORO
11.1. Fica eleito o foro da {foro}, com renúncia a qualquer outro, para dirimir controvérsias decorrentes deste contrato.

E, por estarem de acordo, as partes firmam o presente instrumento.

Local e data: ______________________________

CONTRATANTE: ______________________________

CONTRATADO(A): ______________________________
"""

# =====================================================
# DOCX DOWNLOAD
# =====================================================
def _sanitize_filename(name: str) -> str:
    keep = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_ "
    cleaned = "".join([c if c in keep else "_" for c in (name or "")]).strip()
    return cleaned[:80] if cleaned else "documento"

def _make_docx_bytes(title: str, text: str) -> BytesIO:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for line in (text or "").replace("\r\n", "\n").split("\n"):
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@app.route("/download-docx", methods=["POST"])
def download_docx():
    title = (request.form.get("doc_title") or "Documento").strip()
    text = request.form.get("doc_text") or ""
    filename = _sanitize_filename(request.form.get("doc_filename") or title)

    if not text.strip():
        flash("Nada para baixar. Gere o documento primeiro.", "success")
        return redirect(request.referrer or url_for("home"))

    bio = _make_docx_bytes(title=title, text=text)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"{filename}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =====================================================
# ROTAS
# =====================================================
@app.route("/", methods=["GET", "POST"])
def home():
    answer = None
    if request.method == "POST":
        q = (request.form.get("q") or "").strip()
        if q:
            answer = generate_answer_for_question(q)
            save_history(q, answer)

    return render_template(
        "home.html",
        app_name=APP_NAME,
        history=get_history(50),
        answer=answer,
        questions=QUICK_QUESTIONS,
    )

@app.route("/qa", methods=["GET"])
def qa_get():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"ok": False, "error": "missing q"}), 400
    html = generate_answer_for_question(q)
    return jsonify({"ok": True, "question": q, "answer_html": html})

@app.route("/recursos")
def recursos():
    return render_template("resources.html", app_name=APP_NAME, links=LINKS_OFICIAIS)

@app.route("/contrato", methods=["GET", "POST"])
def contrato():
    contrato_txt = None
    if request.method == "POST":
        contrato_txt = gerar_contrato_advocacia(request.form)
    return render_template("contrato.html", app_name=APP_NAME, contrato_txt=contrato_txt)

# =====================================================
# INIT
# =====================================================
if __name__ == "__main__":
    init_db()
    # Render/produção usa PORT no ambiente
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port)
